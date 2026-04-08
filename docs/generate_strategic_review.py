"""
Generate Lewhofmeyr AI Strategic Review & Roadmap - Professional DOCX
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style Definitions ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level, (size, color) in enumerate([
    (Pt(26), RGBColor(0x1B, 0x2A, 0x4A)),  # Heading 1
    (Pt(18), RGBColor(0x2E, 0x75, 0xB6)),  # Heading 2
    (Pt(14), RGBColor(0x2E, 0x75, 0xB6)),  # Heading 3
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = size
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level > 1 else 24)
    h.paragraph_format.space_after = Pt(8)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1B2A4A"):
    """Create a professional styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_color}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F5F7FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    # Cell margins
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                '  <w:top w:w="60" w:type="dxa"/>'
                '  <w:start w:w="80" w:type="dxa"/>'
                '  <w:bottom w:w="60" w:type="dxa"/>'
                '  <w:end w:w="80" w:type="dxa"/>'
                '</w:tcMar>'
            )
            tcPr.append(tcMar)

    doc.add_paragraph()  # spacing after table
    return table


def add_bullet(doc, text, bold_prefix=None, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_code_block(doc, text):
    """Add a code/prompt block with background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Add shading
    rPr = run._r.get_or_add_rPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F2F5" w:val="clear"/>')
    rPr.append(shading)
    return p


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

# Top spacing
for _ in range(6):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LEWHOFMEYR AI')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
run.font.name = 'Calibri'
run.bold = True

# Divider line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 60)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
run.font.size = Pt(12)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Comprehensive Strategic Review & Roadmap')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Personal AI Operating System | PRD, Technical Spec & Phased Roadmap')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'
run.italic = True

# Spacing
for _ in range(4):
    doc.add_paragraph()

# Meta info
meta_items = [
    ('Prepared For:', 'Lew (COO / Entrepreneur)'),
    ('Prepared By:', 'CTO Office (Claude)'),
    ('Date:', '2026-04-08'),
    ('Classification:', 'Strategic Advisory \u2014 CDDP'),
    ('Version:', '1.0'),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + '  ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Page break
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Executive Summary',
    'Part 1: Current App Classification & Positioning',
    'Part 2: Vision & Future-State Research',
    'Part 3: Product Requirements Document (PRD)',
    'Part 4: Technical Specification',
    'Part 5: Iterative Roadmap',
    'Part 6: Developer Consultant Guide \u2014 Prompting Strategy',
    'Part 7: Cost Projection',
    'Recommendation & Next Steps',
]
for i, item in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}.  {item}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════

doc.add_heading('Executive Summary', level=1)

p = doc.add_paragraph()
p.add_run('Lewhof AI').bold = True
p.add_run(' is a personal AI operating system for high-performing operators and entrepreneurs. It unifies AI-powered chat, task management, calendar, email triage, document intelligence, encrypted credential storage, knowledge base, workflow automation, and autonomous agent orchestration into a single, self-hosted command centre.')

p = doc.add_paragraph()
p.add_run('This document provides a comprehensive strategic review of the current state, gap analysis, product requirements, technical specification, phased roadmap, and developer consultant guidance for evolving the platform into a world-class ')
p.add_run('AI Chief of Staff').bold = True
p.add_run('.')

p = doc.add_paragraph()
p.add_run('Key findings:')
add_bullet(doc, 'The app already has 25+ pages, 76 API routes, and multi-model AI integration \u2014 a strong foundation')
add_bullet(doc, 'Primary gaps: cross-module intelligence, proactive AI, daily planning, finance/health/CRM modules')
add_bullet(doc, 'The 5-phase roadmap delivers quick wins (Phase 0: 1-2 weeks) scaling to world-class (Phase 5: Week 17+)')
add_bullet(doc, 'Total infrastructure cost remains $10-15/mo throughout all phases')
add_bullet(doc, 'No competitor offers this breadth in a single, self-owned platform')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PART 1: CLASSIFICATION & POSITIONING
# ═══════════════════════════════════════════════════════════

doc.add_heading('Part 1: Current App Classification & Positioning', level=1)

doc.add_heading('Classification', level=2)
p = doc.add_paragraph()
run = p.add_run('Personal AI Operating System (Personal AI OS) / Executive Command Centre')
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run('This sits in the emerging category that analysts call "Horizontal AI" \u2014 tools that increase productivity across all functions rather than serving a single vertical. The closest commercial parallels are ')
p.add_run('alfred_').bold = True
p.add_run(' ($25/mo AI chief of staff), ')
p.add_run('Ambient').bold = True
p.add_run(' ($100/mo executive briefing tool), and ')
p.add_run('Motion').bold = True
p.add_run(' ($19/mo AI scheduling) \u2014 but none of them combine all three into one self-owned platform.')

doc.add_heading('Product Description', level=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
p.paragraph_format.right_indent = Inches(0.3)
run = p.add_run('Lewhof AI is a personal AI operating system for high-performing operators and entrepreneurs. It unifies AI-powered chat, task management, calendar, email triage, document intelligence, encrypted credential storage, knowledge base, workflow automation, and autonomous agent orchestration into a single, self-hosted command centre. Built on a dark-native interface with multi-model AI (Claude, Gemini, Groq, Perplexity), it replaces the fragmented stack of productivity apps with one intelligent workspace that learns how you work and acts on your behalf \u2014 giving you back the hours a human chief of staff would cost $150K+/year to provide.')
run.italic = True

doc.add_heading('Current App Inventory', level=2)
inventory = [
    '25+ pages, 76 API routes, 34 components',
    'AI Chat (multi-model: Claude, Gemini, Groq, Perplexity)',
    'Cerebro Master Agent with 20+ tools and slash commands',
    'Custom agent builder (scheduled/event-triggered)',
    'Task management (Kanban), Calendar, Email (Outlook), Notes, Documents',
    'Diagram editor (ReactFlow), Image generation, Workflow builder',
    'Encrypted Vault, Knowledge Base, Whiteboard/backlog',
    'Spotify widget, Weather, Focus timer, Voice input',
    'PWA with offline support, push notifications',
    'Polished dark UI (Warm Dusk theme), responsive mobile',
]
for item in inventory:
    add_bullet(doc, item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PART 2: VISION & FUTURE-STATE RESEARCH
# ═══════════════════════════════════════════════════════════

doc.add_heading('Part 2: Vision & Future-State Research', level=1)

doc.add_heading('Competitive Landscape (April 2026)', level=2)
add_styled_table(doc,
    ['Product', 'What It Does', 'Price', 'Gap vs. Lewhof AI'],
    [
        ['alfred_', 'Email triage + calendar + daily brief', '$25/mo', 'No tasks, no vault, no KB, no agents'],
        ['Ambient', 'Meeting capture + briefings + follow-ups', '$100/mo', 'No personal modules (health, finance)'],
        ['Motion', 'AI scheduling + task auto-placement', '$19/mo', 'No email, no documents, no AI chat'],
        ['Notion', 'Workspace + databases + Notion AI', '$10-20/mo', 'No vault, no email, no agent system'],
        ['Obsidian', 'Local-first PKM + plugins', 'Free-$5/mo', 'No calendar, no email, no AI agents'],
        ['Mem.ai', 'Zero-org AI notes', 'Paid', 'No tasks, no calendar, no automation'],
        ['Lindy.ai', 'No-code agent builder', '$50/mo', 'No personal workspace, agent-only'],
        ['Lewhof AI', 'All of the above, self-owned', '~$10/mo', 'Gaps identified below'],
    ],
    col_widths=[1.2, 2.2, 0.9, 2.5]
)

doc.add_heading('Gap Analysis \u2014 Current vs. World-Class', level=2)
add_styled_table(doc,
    ['Category', 'Current State', 'Gap to World-Class', 'Priority'],
    [
        ['Cross-module intelligence', 'Modules are siloed', 'Calendar should inform tasks, tasks link to notes/KB', 'Critical'],
        ['Proactive AI', 'Cerebro is reactive', 'Should anticipate: overdue follow-ups, at-risk deadlines', 'Critical'],
        ['Daily planning ritual', 'Dashboard shows data only', 'Motion-style "here\'s your optimized day" each morning', 'High'],
        ['Voice interface', 'Web Speech API (basic)', 'Full conversational voice mode, hands-free', 'High'],
        ['Tone/voice learning', 'AI output is generic', 'Learn writing style from history', 'High'],
        ['Finance tracking', 'None', 'Manual entry + CSV import, spending insights', 'Medium'],
        ['CRM/contacts', 'None', 'Auto-extract from email/calendar, relationship health', 'Medium'],
        ['Habit tracking', 'Focus timer only', 'Recurring habits with streaks, linked to calendar', 'Medium'],
        ['Goal/OKR tracking', 'None', 'Quarterly objectives linked to tasks + weekly reviews', 'Medium'],
        ['Health/wellness', 'None', 'Wearable sync, sleep/exercise, mood tracking', 'Medium'],
        ['Web clipper', 'None', 'Bookmarklet to capture pages into KB with AI summary', 'Low'],
        ['MCP integrations', 'None', 'Connect to Google, Slack, financial tools via MCP', 'Low'],
    ],
    col_widths=[1.4, 1.5, 2.3, 0.8]
)

doc.add_heading('The Ultimate Vision \u2014 The AI Chief of Staff', level=2)
p = doc.add_paragraph('A system that:')
vision_items = [
    ('Knows everything ', 'about your schedule, tasks, finances, health, contacts, and goals'),
    ('Connects the dots ', 'across modules (meeting prep pulls relevant notes + contact history + task context)'),
    ('Acts autonomously ', 'within guardrails (drafts emails in your voice, schedules follow-ups, files expenses)'),
    ('Anticipates needs ', '(proactive nudges, pattern detection, risk flagging)'),
    ('Costs $10-15/mo ', 'instead of $150K/year for a human chief of staff'),
]
for bold_part, rest in vision_items:
    add_bullet(doc, rest, bold_prefix=bold_part)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PART 3: PRD
# ═══════════════════════════════════════════════════════════

doc.add_heading('Part 3: Product Requirements Document (PRD)', level=1)

doc.add_heading('Product Vision', level=2)
p = doc.add_paragraph('Transform Lewhof AI from a feature-rich personal dashboard into an intelligent, proactive AI Chief of Staff that manages and optimizes all aspects of a COO/entrepreneur\'s professional and personal life.')

doc.add_heading('Target User', level=2)
add_bullet(doc, 'Lew (COO, entrepreneur, systems thinker)', bold_prefix='Primary: ')
add_bullet(doc, 'High-performing operators, founders, and executives who want one self-owned system instead of 10 SaaS subscriptions', bold_prefix='Future: ')

doc.add_heading('P0 \u2014 Must Have (Phase 1-2)', level=2)
add_styled_table(doc,
    ['ID', 'User Story', 'Acceptance Criteria'],
    [
        ['US-01', 'As a COO, I want a daily AI briefing pushed to me every morning so I start my day informed', 'Briefing includes: weather, calendar, overdue tasks, unread email count, one insight. Delivered by 6:30am SAST via dashboard + Telegram'],
        ['US-02', 'As a COO, I want tasks linked to notes and KB entries so I have context when working on them', 'Task detail view shows linked notes/KB. Notes can embed task references. Bidirectional links'],
        ['US-03', 'As a COO, I want AI to propose my optimal day when I open the planner so I don\'t waste time deciding', 'Smart Planner page: pulls calendar + tasks + habits, proposes time blocks, drag to reorder, click "Lock Day"'],
        ['US-04', 'As a COO, I want Cerebro to proactively flag overdue items and at-risk deadlines', 'Proactive nudge system: checks daily for overdue tasks, stale follow-ups, approaching deadlines. Surfaces on dashboard'],
        ['US-05', 'As a COO, I want quick-capture from anywhere so I can dump ideas without losing flow', 'Enhanced Cmd+K: type anything, AI classifies it (task/note/KB/calendar/whiteboard) and routes automatically'],
    ],
    col_widths=[0.6, 2.8, 3.4]
)

doc.add_heading('P1 \u2014 Should Have (Phase 3-4)', level=2)
add_styled_table(doc,
    ['ID', 'User Story', 'Acceptance Criteria'],
    [
        ['US-06', 'As a COO, I want AI to draft emails in my voice so I can review and send instead of writing from scratch', 'Tone learning from last 50+ messages. Draft quality: <20% editing needed'],
        ['US-07', 'As a COO, I want to track spending categories so I see where my money goes', 'Finance page: manual entry + CSV import. Monthly bar chart by category. AI insight at top'],
        ['US-08', 'As a COO, I want weekly review automation so I reflect on what worked and what didn\'t', 'Weekly review agent: tasks completed, tasks slipped, time allocation, one recommendation'],
        ['US-09', 'As a COO, I want a contacts/CRM layer so I track relationships and never drop follow-ups', 'Auto-extract contacts from email/calendar. Interaction history. "Haven\'t contacted X in 30 days" alerts'],
        ['US-10', 'As a COO, I want recurring habits with streaks so I build consistency', 'Habit tracker: daily/weekly habits, streak count, linked to calendar time blocks, dashboard widget'],
    ],
    col_widths=[0.6, 2.8, 3.4]
)

doc.add_heading('P2 \u2014 Nice to Have (Phase 5+)', level=2)
add_styled_table(doc,
    ['ID', 'User Story', 'Acceptance Criteria'],
    [
        ['US-11', 'As a COO, I want voice-first interaction so I can use the app hands-free', 'Full conversational voice mode with context awareness'],
        ['US-12', 'As a COO, I want goal/OKR tracking so I stay aligned to quarterly objectives', 'Goals page: quarterly OKRs, linked to tasks, AI progress tracking'],
        ['US-13', 'As a COO, I want MCP integrations so external tools feed data into my OS automatically', 'MCP server connections: Google Calendar, Gmail, Slack, financial feeds'],
        ['US-14', 'As a COO, I want a web clipper so I can save any page to KB with one click', 'Browser bookmarklet: captures page, AI generates summary, saves to KB'],
        ['US-15', 'As a COO, I want a journal/reflection module so I track personal growth over time', 'Daily prompts, mood tracking, AI-detected patterns over weeks/months'],
    ],
    col_widths=[0.6, 2.8, 3.4]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PART 4: TECHNICAL SPECIFICATION
# ═══════════════════════════════════════════════════════════

doc.add_heading('Part 4: Technical Specification', level=1)

doc.add_heading('4.1 Current Architecture', level=2)
arch_layers = [
    ('Frontend:', 'Next.js 16.2.2 (App Router, RSC), Tailwind v4 + Warm Dusk Design System, 25+ pages, 34 components, PWA'),
    ('API Layer:', '76 API Routes (Next.js Route Handlers), Clerk Auth (Google/Microsoft/Spotify OAuth), Anthropic SDK, Gemini, Groq, Perplexity'),
    ('Data Layer:', 'Supabase (Postgres + Storage + RLS), AES Encryption (Vault), IndexedDB (Offline/PWA)'),
    ('Integrations:', 'Microsoft Graph (Calendar/Email), Spotify, Telegram, GitHub, Vercel (Hosting + Cron)'),
]
for label, desc in arch_layers:
    add_bullet(doc, desc, bold_prefix=label + ' ')

doc.add_heading('4.2 Proposed Architecture Additions', level=2)

p = doc.add_paragraph()
run = p.add_run('Phase 1-2 \u2014 No new infrastructure')
run.bold = True
add_styled_table(doc,
    ['Addition', 'Implementation', 'Cost Impact'],
    [
        ['Cross-module links', 'New entity_links Supabase table', '$0'],
        ['Daily briefing agent', 'Vercel Cron \u2192 existing APIs \u2192 Telegram', '$0'],
        ['Smart planner page', 'New page consuming existing calendar + todos APIs', '$0'],
        ['Proactive nudge engine', 'Vercel Cron job running daily checks against task deadlines', '$0'],
        ['Enhanced quick-capture', 'AI classification layer in existing Cmd+K modal', '~$0.01/classify'],
    ],
    col_widths=[1.8, 3.2, 1.2]
)

p = doc.add_paragraph()
run = p.add_run('Phase 3-4 \u2014 Minimal new infrastructure')
run.bold = True
add_styled_table(doc,
    ['Addition', 'Implementation', 'Cost Impact'],
    [
        ['Tone learning', 'New user_writing_profile table, batch analysis', '~$0.50 one-time'],
        ['Finance tracker', 'New finance_entries table + page', '$0'],
        ['CRM/contacts', 'New contacts table, auto-extract from email/calendar', '$0'],
        ['Habit tracker', 'New habits + habit_logs tables + dashboard widget', '$0'],
        ['Weekly review agent', 'Scheduled Vercel Cron, queries existing data', '~$0.05/week'],
    ],
    col_widths=[1.8, 3.2, 1.2]
)

p = doc.add_paragraph()
run = p.add_run('Phase 5+ \u2014 Strategic investments')
run.bold = True
add_styled_table(doc,
    ['Addition', 'Implementation', 'Cost Impact'],
    [
        ['MCP integration', 'MCP client in API routes, external MCP servers', '$0 (protocol free)'],
        ['Voice mode', 'Web Speech API upgrade or Whisper API', '~$0.006/min'],
        ['Model routing', 'LiteLLM proxy or OpenRouter', 'Saves 40-60%'],
        ['Semantic caching', 'Redis/Upstash for repeated AI queries', 'Free \u2192 $10/mo'],
    ],
    col_widths=[1.8, 3.2, 1.2]
)

doc.add_heading('4.3 Data Model Additions', level=2)

tables_spec = [
    ('entity_links (Phase 1)', [
        ['id', 'UUID', 'Primary Key'],
        ['user_id', 'TEXT', 'NOT NULL'],
        ['source_type', 'TEXT', 'task / note / kb / calendar / contact'],
        ['source_id', 'UUID', 'NOT NULL'],
        ['target_type', 'TEXT', 'NOT NULL'],
        ['target_id', 'UUID', 'NOT NULL'],
        ['created_at', 'TIMESTAMPTZ', 'DEFAULT now()'],
    ]),
    ('daily_plans (Phase 2)', [
        ['id', 'UUID', 'Primary Key'],
        ['user_id', 'TEXT', 'NOT NULL'],
        ['plan_date', 'DATE', 'NOT NULL'],
        ['blocks', 'JSONB', 'Array of {time, type, ref_id, title, duration}'],
        ['locked', 'BOOLEAN', 'DEFAULT false'],
        ['created_at', 'TIMESTAMPTZ', 'DEFAULT now()'],
    ]),
    ('finance_entries (Phase 3)', [
        ['id', 'UUID', 'Primary Key'],
        ['user_id', 'TEXT', 'NOT NULL'],
        ['amount', 'DECIMAL', 'NOT NULL'],
        ['category', 'TEXT', 'NOT NULL'],
        ['description', 'TEXT', ''],
        ['entry_date', 'DATE', 'NOT NULL'],
        ['type', 'TEXT', 'expense / income'],
        ['created_at', 'TIMESTAMPTZ', 'DEFAULT now()'],
    ]),
    ('contacts (Phase 3)', [
        ['id', 'UUID', 'Primary Key'],
        ['user_id', 'TEXT', 'NOT NULL'],
        ['name', 'TEXT', 'NOT NULL'],
        ['email', 'TEXT', ''],
        ['phone', 'TEXT', ''],
        ['company', 'TEXT', ''],
        ['tags', 'TEXT[]', 'Array'],
        ['last_contact', 'TIMESTAMPTZ', ''],
        ['notes', 'TEXT', ''],
        ['source', 'TEXT', 'email / calendar / manual'],
        ['created_at', 'TIMESTAMPTZ', 'DEFAULT now()'],
    ]),
    ('habits (Phase 3)', [
        ['id', 'UUID', 'Primary Key'],
        ['user_id', 'TEXT', 'NOT NULL'],
        ['name', 'TEXT', 'NOT NULL'],
        ['frequency', 'TEXT', 'daily / weekly'],
        ['target_days', 'INTEGER', 'DEFAULT 1'],
        ['current_streak', 'INTEGER', 'DEFAULT 0'],
        ['best_streak', 'INTEGER', 'DEFAULT 0'],
        ['active', 'BOOLEAN', 'DEFAULT true'],
        ['created_at', 'TIMESTAMPTZ', 'DEFAULT now()'],
    ]),
    ('habit_logs (Phase 3)', [
        ['id', 'UUID', 'Primary Key'],
        ['habit_id', 'UUID', 'FK \u2192 habits(id)'],
        ['completed_at', 'TIMESTAMPTZ', 'DEFAULT now()'],
    ]),
]

for table_name, columns in tables_spec:
    p = doc.add_paragraph()
    run = p.add_run(table_name)
    run.bold = True
    run.font.size = Pt(11)
    add_styled_table(doc,
        ['Column', 'Type', 'Notes'],
        columns,
        col_widths=[1.5, 1.5, 3.5]
    )

doc.add_heading('4.4 Security Considerations', level=2)
add_styled_table(doc,
    ['Concern', 'Current State', 'Recommendation'],
    [
        ['Data encryption', 'Vault uses AES', 'Extend to finance entries, contacts'],
        ['API auth', 'Clerk + RLS', 'Sufficient for personal use'],
        ['Token exposure', 'API keys in env vars', 'Vercel encrypted env vars (already there)'],
        ['MCP security', 'Not yet implemented', 'Authenticate every MCP connection, no public servers'],
        ['Finance data', 'Not yet stored', 'PIN-gate finance page like Vault'],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_heading('4.5 Scalability', level=2)
p = doc.add_paragraph('Current architecture (Vercel serverless + Supabase) scales to 10,000+ users without changes. For personal use, zero scalability concerns for 2+ years. Commercial upgrade path: Supabase Pro ($25/mo) at 500MB DB, Vercel Pro ($20/mo) at execution limits.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PART 5: ITERATIVE ROADMAP
# ═══════════════════════════════════════════════════════════

doc.add_heading('Part 5: Iterative Roadmap', level=1)

phases = [
    {
        'title': 'Phase 0 \u2014 Foundation Hardening (Week 1-2)',
        'subtitle': 'Quick wins that make the existing app feel alive.',
        'features': [
            ['0.1', 'Daily briefing agent (Vercel Cron \u2192 Telegram)', 'Small', 'High', 'Existing APIs'],
            ['0.2', 'Cross-module entity linking (tasks \u2194 notes \u2194 KB)', 'Medium', 'High', 'New table'],
            ['0.3', 'Whiteboard auto-prioritization (AI scoring)', 'Small', 'Medium', 'Existing whiteboard'],
            ['0.4', 'Enhanced Cmd+K quick-capture with AI routing', 'Medium', 'High', 'Existing search modal'],
            ['0.5', 'Email triage polish (auto-labels, summaries)', 'Small', 'Medium', 'Existing email integration'],
        ],
        'metric': 'Daily briefing arrives by 6:30am SAST for 7 consecutive days.',
    },
    {
        'title': 'Phase 1 \u2014 Intelligent Planning (Week 3-5)',
        'subtitle': 'The single biggest UX leap \u2014 your day planned for you.',
        'features': [
            ['1.1', 'Smart Daily Planner page', 'Large', 'Critical', 'Calendar + todos APIs'],
            ['1.2', 'Task auto-scheduling (deadline + priority \u2192 time block)', 'Medium', 'High', 'Planner page'],
            ['1.3', 'Focus time auto-blocking on calendar', 'Small', 'Medium', 'Calendar write API'],
            ['1.4', 'Weekly review agent (end-of-week summary)', 'Medium', 'High', 'Cron + existing data'],
            ['1.5', 'Habit tracker (recurring habits, streaks, widget)', 'Medium', 'Medium', 'New tables'],
        ],
        'metric': 'You use the Smart Planner every morning for 14 consecutive days.',
    },
    {
        'title': 'Phase 2 \u2014 AI That Knows You (Week 6-8)',
        'subtitle': 'The moat \u2014 personalized, proactive intelligence.',
        'features': [
            ['2.1', 'Proactive nudge engine (overdue, at-risk, follow-ups)', 'Medium', 'Critical', 'Cron + todos + email'],
            ['2.2', 'Tone/voice learning (AI writes like you)', 'Medium', 'High', 'Historical message analysis'],
            ['2.3', 'AI email draft generation (in your voice)', 'Medium', 'High', 'Tone profile + email API'],
            ['2.4', 'Semantic search across all modules', 'Medium', 'High', 'Embedding generation'],
            ['2.5', 'Contact/CRM auto-extraction from email + calendar', 'Medium', 'Medium', 'New table + Graph API'],
        ],
        'metric': 'AI drafts need <20% editing. At least one proactive nudge saves you from missing something.',
    },
    {
        'title': 'Phase 3 \u2014 Life Modules (Week 9-12)',
        'subtitle': 'Expand from work productivity to full life OS.',
        'features': [
            ['3.1', 'Finance tracker (manual + CSV import + insights)', 'Medium', 'High', 'New table'],
            ['3.2', 'Goal/OKR tracker (quarterly objectives linked to tasks)', 'Medium', 'Medium', 'New table + todos'],
            ['3.3', 'Web clipper bookmarklet (save pages to KB)', 'Small', 'Medium', 'KB API'],
            ['3.4', 'Journal/reflection module (prompts, mood, AI patterns)', 'Medium', 'Medium', 'New table'],
            ['3.5', 'Relationship health dashboard (contact frequency)', 'Small', 'Medium', 'CRM data'],
        ],
        'metric': 'Full month of finances tracked. Quarterly OKRs set and linked to weekly tasks.',
    },
    {
        'title': 'Phase 4 \u2014 Automation & Integration (Week 13-16)',
        'subtitle': 'Let the system work while you sleep.',
        'features': [
            ['4.1', 'Model routing (Haiku/Sonnet/Opus by complexity)', 'Medium', 'High', 'LiteLLM or custom router'],
            ['4.2', 'Semantic caching (Redis/Upstash)', 'Medium', 'Medium', 'Cache layer'],
            ['4.3', 'MCP client integration (Google Calendar, Gmail)', 'Large', 'High', 'MCP protocol'],
            ['4.4', 'Agent marketplace (meeting prep, expense logger)', 'Medium', 'High', 'Agent framework'],
            ['4.5', 'Two-way Telegram (send commands back)', 'Small', 'Medium', 'Telegram bot update'],
        ],
        'metric': 'AI costs reduced 40%+. 3+ agents running autonomously.',
    },
    {
        'title': 'Phase 5 \u2014 Polish & Scale (Week 17+)',
        'subtitle': 'World-class finish.',
        'features': [
            ['5.1', 'Full voice mode (conversational, context-aware)', 'Large', 'High', 'Speech APIs'],
            ['5.2', 'Mobile-native build (Capacitor, gesture nav)', 'Large', 'High', 'Capacitor config'],
            ['5.3', 'Onboarding flow (first-run wizard, feature tour)', 'Medium', 'Medium', 'All modules'],
            ['5.4', 'Data export/import (full portability)', 'Medium', 'Medium', 'All tables'],
            ['5.5', 'Multi-user architecture prep (team/family)', 'Large', 'Future', 'Schema updates'],
        ],
        'metric': 'A friend/colleague says "I want this" unprompted.',
    },
]

for phase in phases:
    doc.add_heading(phase['title'], level=2)
    p = doc.add_paragraph(phase['subtitle'])
    p.runs[0].italic = True

    add_styled_table(doc,
        ['#', 'Feature', 'Effort', 'Impact', 'Dependencies'],
        phase['features'],
        col_widths=[0.4, 2.8, 0.8, 0.8, 1.8]
    )

    p = doc.add_paragraph()
    run = p.add_run('Success Metric: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.add_run(phase['metric'])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PART 6: DEVELOPER CONSULTANT GUIDE
# ═══════════════════════════════════════════════════════════

doc.add_heading('Part 6: Developer Consultant Guide \u2014 Prompting Strategy', level=1)

doc.add_heading('The Prompting Framework', level=2)
p = doc.add_paragraph('Every phase execution follows this structure when working with Claude Code:')

framework_items = [
    ('1. CONTEXT', 'What exists, what we\'re building on'),
    ('2. OBJECTIVE', 'Exactly what to build'),
    ('3. CONSTRAINTS', 'Tech stack, patterns, existing conventions'),
    ('4. ACCEPTANCE', 'How to verify it works'),
    ('5. QUALITY', 'QA checks to run'),
]
for label, desc in framework_items:
    add_bullet(doc, desc, bold_prefix=label + ' \u2014 ')

doc.add_heading('Phase 0 Prompt Templates', level=2)

doc.add_heading('0.1 \u2014 Daily Briefing Agent', level=3)
add_code_block(doc, 'Build a daily briefing system.\n\nCONTEXT: We have existing API routes at /api/weather, /api/calendar/events, /api/todos/pending, /api/email. Telegram bot is active at /api/telegram.\n\nOBJECTIVE: Create a Vercel Cron job that runs at 04:30 UTC (06:30 SAST) daily. It should:\n1. Call each API to gather: weather, today\'s events, overdue tasks, unread email count\n2. Format into a clean briefing message\n3. Send via Telegram bot\n4. Also store in Supabase for dashboard display\n\nCONSTRAINTS: Use existing API routes, don\'t duplicate logic. Vercel Cron config in vercel.json. Match existing code patterns in /api/cron/.\n\nACCEPTANCE CRITERIA: Briefing arrives on Telegram by 06:30 SAST. Dashboard briefing widget shows latest briefing. Handles API failures gracefully.\n\nQA: Run CTO protocol. Verify Vercel build. Test with manual cron trigger first.')

doc.add_heading('0.4 \u2014 Enhanced Quick-Capture', level=3)
add_code_block(doc, 'Enhance the existing Cmd+K global search modal to support intelligent quick-capture.\n\nCONTEXT: GlobalSearch component exists at src/components/global-search.tsx. Currently does keyword search.\n\nOBJECTIVE: When user types in Cmd+K and presses Enter:\n1. AI classifies input as: task, note, kb_entry, calendar_event, or whiteboard_item\n2. Routes to the correct API to create the item\n3. Shows confirmation toast with item type icon\n4. Falls back to search if classification confidence is low\n\nCONSTRAINTS: Use Haiku for classification (cheapest, fastest). Keep existing search functionality intact.\n\nACCEPTANCE CRITERIA: "Buy milk tomorrow" \u2192 creates task with due date. "Meeting with John at 3pm" \u2192 creates calendar event. Ambiguous input \u2192 falls back to search.')

doc.add_heading('Phase 1 Prompt Templates', level=2)

doc.add_heading('1.1 \u2014 Smart Daily Planner', level=3)
add_code_block(doc, 'Build a Smart Daily Planner page at /planner.\n\nCONTEXT: We have /api/calendar/events for today\'s events and /api/todos for tasks with priority and due dates.\n\nOBJECTIVE: New page that on load:\n1. Fetches today\'s calendar events (fixed blocks)\n2. Fetches incomplete tasks sorted by priority + due date\n3. AI proposes time blocks filling gaps between meetings\n4. User sees a vertical timeline view (6am-10pm)\n5. User can drag to reorder suggested blocks\n6. "Lock Day" button commits the plan\n\nCONSTRAINTS: Timeline UI matches Warm Dusk design system. Fixed events shown in muted style. AI-suggested blocks with dashed border. Mobile: stack vertically, tap to reorder.\n\nACCEPTANCE CRITERIA: Page loads in <2s. Calendar events locked. Tasks fill gaps by priority. Drag reorder on desktop. "Lock Day" persists.')

doc.add_heading('Quality Assurance Prompt (Use After Every Feature)', level=2)
add_code_block(doc, 'Run full QA on everything just built:\n1. TypeScript check: run tsc --noEmit, show output\n2. Build check: run next build, show any errors\n3. Console.log audit: grep for console.log in changed files\n4. Hardcoded values: check for any hardcoded URLs, keys, or user-specific values\n5. Dead code: check for unused imports and variables\n6. Mobile layout: verify at 375px width\n7. Desktop layout: verify at 1280px width\n8. If new API routes: test with curl, show response\n9. If new DB tables: verify RLS policies are set\n10. Git push and verify Vercel deployment reaches READY')

doc.add_heading('Debugging Prompt', level=2)
add_code_block(doc, '[paste screenshot]\n\nThis is what I see on [device/browser].\n\nExpected behavior: [what should happen]\nActual behavior: [what\'s happening]\n\nInvestigate, identify root cause, fix it, and run QA protocol. Don\'t ask questions \u2014 just fix it.')

doc.add_heading('Validation Checklist Per Phase', level=2)
add_styled_table(doc,
    ['Phase', 'Validation Steps'],
    [
        ['Phase 0', 'Briefing delivered 7 days straight. Cross-links work bidirectionally. Quick-capture correctly routes 5/5 test inputs.'],
        ['Phase 1', 'Planner proposes sensible day for 14 consecutive days. Habits tracked with visible streaks. Weekly review generates useful summary.'],
        ['Phase 2', 'AI drafts match your tone (blind test). At least 3 proactive nudges per week that are useful. Semantic search returns relevant results.'],
        ['Phase 3', 'Full month of finances tracked without gaps. OKRs set and tasks visibly linked. Journal has 20+ entries.'],
        ['Phase 4', 'AI costs down 40%+ (compare monthly spend). 3+ agents running unattended. MCP connects to at least one external service.'],
    ],
    col_widths=[1.0, 5.8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# PART 7: COST PROJECTION
# ═══════════════════════════════════════════════════════════

doc.add_heading('Part 7: Cost Projection', level=1)

add_styled_table(doc,
    ['Phase', 'Timeline', 'Additional Cost', 'Cumulative Monthly'],
    [
        ['Current', '\u2014', '$0', '~$10/mo (API + infra)'],
        ['Phase 0', 'Week 1-2', '$0', '~$10/mo'],
        ['Phase 1', 'Week 3-5', '$0', '~$12/mo'],
        ['Phase 2', 'Week 6-8', '$0', '~$15/mo'],
        ['Phase 3', 'Week 9-12', '$0', '~$15/mo'],
        ['Phase 4', 'Week 13-16', 'Upstash Redis free tier', '~$10/mo (savings from routing)'],
        ['Phase 5', 'Week 17+', 'Capacitor build tools (free)', '~$10-15/mo'],
    ],
    col_widths=[1.2, 1.5, 2.0, 2.0]
)

p = doc.add_paragraph()
run = p.add_run('Total investment over 4 months: $40-60 in AI API costs.')
run.bold = True
run.font.size = Pt(12)
p.add_run(' That\'s less than one month of any competitor subscription.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# RECOMMENDATION
# ═══════════════════════════════════════════════════════════

doc.add_heading('Recommendation & Next Steps', level=1)

p = doc.add_paragraph()
p.add_run('Start with Phase 0.1 (Daily Briefing Agent)').bold = True
p.add_run(' \u2014 it\'s the highest-impact, lowest-effort win. It makes the app feel alive every morning and proves the autonomous agent pattern that underpins everything else in the roadmap.')

doc.add_heading('Immediate Actions', level=2)
actions = [
    'Approve this strategic roadmap and confirm Phase 0 priorities',
    'Begin CDDP on Phase 0.1 (Daily Briefing Agent) \u2014 full technical spec for your approval',
    'Set up Vercel Cron configuration for scheduled agent execution',
    'Establish the entity_links table for cross-module intelligence (Phase 0.2)',
    'Review and adjust timeline estimates based on availability',
]
for i, action in enumerate(actions, 1):
    add_bullet(doc, action, bold_prefix=f'{i}. ')

# Footer note
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2014 End of Document \u2014')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared by CTO Office (Claude) | Lewhof AI | 2026-04-08')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(9)

# ── Save ──
import pathlib
output_path = pathlib.Path(r'C:\Users\admin\my-ai-tool\docs\Lewhofmeyr-AI-Strategic-Review.docx')
output_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(output_path))
print(f'File size: {output_path.stat().st_size / 1024:.1f} KB')
print(f'Document saved to: {output_path}')
